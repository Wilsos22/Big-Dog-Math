from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output" / "docx" / "M1.T3"
OUT.mkdir(parents=True, exist_ok=True)

FONT = "Arial"
INK = "201E1A"
SOFT_INK = "6F675C"
CREAM = "FAF6EE"
WHITE = "FFFFFF"
LINE = "D9D0C0"
GOLD = "D69C2F"


@dataclass(frozen=True)
class Problem:
    number: str
    prompt: str
    labels: tuple[str, ...] = ()
    note: str = ""


@dataclass(frozen=True)
class KeyItem:
    number: str
    answer: str
    evidence: str = ""


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_run_font(run, size: float, bold: bool = False, color: str = INK, italic: bool = False):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = rgb(color)
    return run


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 90, start: int = 110, bottom: int = 90, end: int = 110):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    # Write both logical and physical horizontal margins. Word honors start/end,
    # while LibreOffice's headless PDF renderer relies on left/right.
    for margin, value in (
        ("top", top),
        ("start", start),
        ("left", start),
        ("bottom", bottom),
        ("end", end),
        ("right", end),
    ):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, color: str = LINE, size: int = 8, edges: Sequence[str] = ("top", "left", "bottom", "right")):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in edges:
        tag = "start" if edge == "left" else "end" if edge == "right" else edge
        element = borders.find(qn(f"w:{tag}"))
        if element is None:
            element = OxmlElement(f"w:{tag}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths_in: Sequence[float], indent_dxa: int = 0):
    total_dxa = round(sum(widths_in) * 1440)
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_in:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(round(width * 1440)))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths_in):
            cell.width = Inches(width)
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(round(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def set_row_min_height(row, inches: float):
    tr_pr = row._tr.get_or_add_trPr()
    height = OxmlElement("w:trHeight")
    height.set(qn("w:val"), str(round(inches * 1440)))
    height.set(qn("w:hRule"), "atLeast")
    tr_pr.append(height)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def paragraph(cell_or_doc, text: str = "", *, size: float = 10.5, bold: bool = False,
              color: str = INK, before: float = 0, after: float = 2,
              align=WD_ALIGN_PARAGRAPH.LEFT, italic: bool = False):
    p = cell_or_doc.add_paragraph() if hasattr(cell_or_doc, "add_paragraph") else None
    if p is None:
        raise TypeError("Object cannot accept paragraphs")
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.05
    if text:
        set_run_font(p.add_run(text), size, bold, color, italic)
    return p


def clear_paragraph(paragraph_obj):
    for child in list(paragraph_obj._p):
        paragraph_obj._p.remove(child)


def configure_document(accent: str, teacher: bool = False) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.42)
    section.bottom_margin = Inches(0.42)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.header_distance = Inches(0.2)
    section.footer_distance = Inches(0.22)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.05

    for style_name, size, color in (("Heading 1", 14, accent), ("Heading 2", 11.5, INK), ("Heading 3", 10.5, SOFT_INK)):
        style = styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(6)
        style.paragraph_format.space_after = Pt(3)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(0)
    fp.paragraph_format.space_after = Pt(0)
    set_run_font(fp.add_run("Big Dog Math | M1.T3 Decimals" + (" | Teacher Key" if teacher else "")), 8, True, SOFT_INK)
    return doc


def add_title(doc: Document, code: str, title: str, accent: str, label: str = "STUDENT RESOURCE"):
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, (1.35, 6.15))
    left, right = table.rows[0].cells
    for cell in (left, right):
        set_cell_margins(cell, 90, 130, 90, 130)
        set_cell_border(cell, accent, 10)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(left, accent)
    set_cell_shading(right, CREAM)
    clear_paragraph(left.paragraphs[0])
    left.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(left.paragraphs[0].add_run(code), 13, True, WHITE)
    clear_paragraph(right.paragraphs[0])
    set_run_font(right.paragraphs[0].add_run(title), 18, True, INK)
    rp = right.add_paragraph()
    rp.paragraph_format.space_before = Pt(1)
    rp.paragraph_format.space_after = Pt(0)
    set_run_font(rp.add_run(label), 8.5, True, accent)
    paragraph(doc, "Name: ____________________________________   Period: ______   Date: _______________", size=9.5, bold=True, color=SOFT_INK, before=4, after=4)


def add_directions_band(doc: Document, text: str, accent: str, title: str = "DIRECTIONS"):
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, (1.15, 6.35))
    label_cell, text_cell = table.rows[0].cells
    for cell in (label_cell, text_cell):
        set_cell_margins(cell, 75, 110, 75, 110)
        set_cell_border(cell, LINE, 8)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(label_cell, accent)
    set_cell_shading(text_cell, "FFFDF8")
    clear_paragraph(label_cell.paragraphs[0])
    label_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(label_cell.paragraphs[0].add_run(title), 8.5, True, WHITE)
    clear_paragraph(text_cell.paragraphs[0])
    set_run_font(text_cell.paragraphs[0].add_run(text), 9.5, True, INK)
    paragraph(doc, "", size=1, after=1)


def add_work_card(doc: Document, problem: Problem, accent: str, min_height: float = 2.2):
    table = doc.add_table(rows=2, cols=1)
    set_table_geometry(table, (7.5,))
    head, body = table.rows[0].cells[0], table.rows[1].cells[0]
    prevent_row_split(table.rows[0])
    prevent_row_split(table.rows[1])
    set_cell_shading(head, accent)
    set_cell_shading(body, WHITE)
    for cell in (head, body):
        set_cell_margins(cell, 85, 130, 85, 130)
        set_cell_border(cell, accent if cell is head else LINE, 8)
    clear_paragraph(head.paragraphs[0])
    set_run_font(head.paragraphs[0].add_run(f"{problem.number}. {problem.prompt}"), 10.5, True, WHITE)
    clear_paragraph(body.paragraphs[0])
    body.paragraphs[0].paragraph_format.space_after = Pt(2)
    if problem.note:
        set_run_font(body.paragraphs[0].add_run(problem.note), 9, False, SOFT_INK, True)
    for label in problem.labels:
        p = body.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(1)
        set_run_font(p.add_run(label), 8.5, True, accent)
        line = body.add_paragraph()
        line.paragraph_format.space_after = Pt(3)
        line.paragraph_format.line_spacing = 1.0
        set_run_font(line.add_run("________________________________________________________________________________"), 8.5, False, LINE)
    if not problem.labels:
        for _ in range(4):
            line = body.add_paragraph()
            line.paragraph_format.space_after = Pt(2)
            set_run_font(line.add_run("________________________________________________________________________________"), 8.5, False, LINE)
    set_row_min_height(table.rows[1], min_height)
    paragraph(doc, "", size=1, after=2)


def add_challenge(doc: Document, text: str, lines: int = 3):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, (7.5,))
    cell = table.cell(0, 0)
    set_cell_margins(cell, 100, 140, 100, 140)
    set_cell_border(cell, GOLD, 12)
    set_cell_shading(cell, "FFF7E2")
    clear_paragraph(cell.paragraphs[0])
    set_run_font(cell.paragraphs[0].add_run("OPTIONAL BIG DOG CHALLENGE"), 9, True, "7A5600")
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    set_run_font(p.add_run(text), 10, True, INK)
    for _ in range(lines):
        line = cell.add_paragraph()
        line.paragraph_format.space_after = Pt(2)
        set_run_font(line.add_run("________________________________________________________________________________"), 8.5, False, "E7D29D")


def add_answer_table(doc: Document, items: Sequence[KeyItem], accent: str, widths=(0.55, 3.2, 3.75)):
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, widths)
    headers = ("#", "Answer and work", "Instructional evidence")
    for cell, text in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, accent)
        set_cell_border(cell, accent, 8)
        set_cell_margins(cell, 75, 95, 75, 95)
        clear_paragraph(cell.paragraphs[0])
        set_run_font(cell.paragraphs[0].add_run(text), 9, True, WHITE)
    set_repeat_table_header(table.rows[0])
    for item in items:
        row = table.add_row()
        prevent_row_split(row)
        for cell in row.cells:
            set_cell_border(cell, LINE, 7)
            set_cell_margins(cell, 85, 95, 85, 95)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        clear_paragraph(row.cells[0].paragraphs[0])
        row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(row.cells[0].paragraphs[0].add_run(item.number), 9.5, True, accent)
        clear_paragraph(row.cells[1].paragraphs[0])
        set_run_font(row.cells[1].paragraphs[0].add_run(item.answer), 9.2, False, INK)
        clear_paragraph(row.cells[2].paragraphs[0])
        set_run_font(row.cells[2].paragraphs[0].add_run(item.evidence), 8.8, False, SOFT_INK)
    return table


def save(doc: Document, filename: str):
    path = OUT / filename
    doc.save(path)
    return path


def build_division_packet(code: str, title: str, accent: str, directions: str,
                          problems: Sequence[Problem], challenge: str,
                          key_items: Sequence[KeyItem], key_challenge: str):
    doc = configure_document(accent)
    add_title(doc, code, title, accent)
    add_directions_band(doc, directions, accent)
    for index, problem in enumerate(problems):
        if index == 2:
            doc.add_page_break()
            add_title(doc, code, title, accent)
        add_work_card(doc, problem, accent, min_height=2.22 if index < 2 else 2.0)
    add_challenge(doc, challenge)
    student_path = save(doc, f"{code}-{title.replace(' ', '-')}-Student.docx")

    key = configure_document(accent, teacher=True)
    add_title(key, code, title, accent, label="TEACHER KEY")
    add_directions_band(key, "Use the first incorrect line to choose the next support. Every student completes all four required items; the challenge is separate.", accent, title="KEY USE")
    add_answer_table(key, key_items, accent)
    paragraph(key, "Challenge", size=10, bold=True, color="7A5600", before=6, after=2)
    paragraph(key, key_challenge, size=9.5, after=3)
    key_path = save(key, f"{code}-{title.replace(' ', '-')}-Teacher-Key.docx")
    return student_path, key_path


def build_reference():
    accent = "6A5743"
    doc = configure_document(accent)
    add_title(doc, "M1.T3", "Decimal Operations Reference", accent, label="COMMON ONE-PAGE REFERENCE")
    add_directions_band(doc, "Use this page to remember the structure. It supports your thinking; it does not replace showing your work.", accent, title="HOW TO USE")

    cards = [
        ("COMPARE AND ORDER", "Align place values. Add placeholder zeros when useful. Compare from left to right. On a number line, find the endpoints and one interval before placing a point."),
        ("ADD AND SUBTRACT", "Estimate. Align decimal points and like places. Add placeholder zeros. Regroup by place value. Compute. Check with the inverse operation and the estimate."),
        ("MULTIPLY", "Estimate the product's size. Use an area model or partial products. Multiply. Place the decimal using place value and magnitude. Check whether the result fits the estimate."),
        ("WHOLE-NUMBER DIVISION", "Estimate. Place each quotient digit by its value. Divide, multiply, subtract, rename and combine, then repeat. Verify: divisor x quotient + remainder = dividend."),
        ("DECIMAL DIVISOR", "Make the divisor whole by multiplying both the divisor and dividend by the same power of 10. Write the equivalent expression. Divide. Verify with the original divisor."),
        ("REMAINDERS", "A remainder must be smaller than the divisor. The question decides whether to report it, convert it, discard it, or round up."),
    ]
    table = doc.add_table(rows=3, cols=2)
    set_table_geometry(table, (3.7, 3.8))
    for cell, (heading, body) in zip((c for row in table.rows for c in row.cells), cards):
        set_cell_border(cell, LINE, 8)
        set_cell_margins(cell, 105, 125, 105, 125)
        set_cell_shading(cell, "FFFDF8")
        clear_paragraph(cell.paragraphs[0])
        set_run_font(cell.paragraphs[0].add_run(heading), 9.2, True, accent)
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.02
        set_run_font(p.add_run(body), 9.2, False, INK)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    for row in table.rows:
        prevent_row_split(row)
        set_row_min_height(row, 1.35)

    paragraph(doc, "TASK ATTACK", size=10, bold=True, color=accent, before=6, after=2)
    task_table = doc.add_table(rows=1, cols=7)
    set_table_geometry(task_table, (1.05, 1.0, 0.95, 1.15, 1.15, 1.05, 1.15))
    steps = (
        "1 Underline important information",
        "2 Cross out what does not help",
        "3 Circle what is asked",
        "4 Rewrite the useful facts",
        "5 List possible strategies",
        "6 Try one visible first step",
        "7 Check the result in context",
    )
    for index, (cell, text) in enumerate(zip(task_table.rows[0].cells, steps)):
        set_cell_border(cell, accent, 8)
        set_cell_margins(cell, 75, 75, 75, 75)
        set_cell_shading(cell, CREAM if index % 2 == 0 else "FFFDF8")
        clear_paragraph(cell.paragraphs[0])
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(cell.paragraphs[0].add_run(text), 7.6, True, INK)
    paragraph(doc, "A strong check answers two questions: Is the computation correct? Does the answer make sense here?", size=9.2, bold=True, color=SOFT_INK, before=4, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)
    return save(doc, "M1.T3-Decimal-Operations-Reference.docx")


def add_review_grid(doc: Document, problems: Sequence[Problem], accent: str):
    table = doc.add_table(rows=3, cols=2)
    set_table_geometry(table, (3.75, 3.75))
    for cell, problem in zip((c for row in table.rows for c in row.cells), problems):
        set_cell_border(cell, LINE, 8)
        set_cell_margins(cell, 90, 110, 90, 110)
        set_cell_shading(cell, WHITE)
        clear_paragraph(cell.paragraphs[0])
        set_run_font(cell.paragraphs[0].add_run(f"{problem.number}. {problem.prompt}"), 9.3, True, INK)
        for _ in range(3):
            p = cell.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            set_run_font(p.add_run("____________________________________"), 8, False, LINE)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    for row in table.rows:
        prevent_row_split(row)
        set_row_min_height(row, 2.25)


def build_review():
    accent = "A34F28"
    problems = [
        Problem("1", "Order 0.47, 0.407, 0.74, and 0.704 from least to greatest."),
        Problem("2", "A line from 0.60 to 0.70 is divided into 10 equal intervals. What value is three intervals to the right of 0.60?"),
        Problem("3", "Find 6.35 + 0.728. Show aligned place values and an estimate."),
        Problem("4", "Find 12.8 - 3.47. Show aligned place values and an estimate."),
        Problem("5", "Find 1.8 x 0.4. Use magnitude to explain the decimal placement."),
        Problem("6", "Find 2.4 x 1.3. Show partial products or a labeled area model."),
        Problem("7", "Find 2,436 ÷ 4. Explain the zero in the quotient."),
        Problem("8", "Find 7.2 ÷ 0.3. Write the equivalent whole-divisor expression."),
        Problem("9", "Repair: 4.7 + 0.386 = 0.856 because the final digits were aligned."),
        Problem("10", "A rectangle measures 8.4 m by 2.5 m. Find and label its area."),
        Problem("11", "193 students ride buses that hold 24 students. How many buses are needed?"),
        Problem("12", "Explain why 6.4 ÷ 0.8 and 64 ÷ 8 have the same quotient. Find it."),
    ]
    doc = configure_document(accent)
    add_title(doc, "M1.T3", "BRUH Review Guide", accent, label="REQUIRED PAPER REVIEW")
    add_directions_band(doc, "Complete every problem. Show enough work to defend each answer. Game points do not replace the paper evidence.", accent)
    add_review_grid(doc, problems[:6], accent)
    doc.add_page_break()
    add_title(doc, "M1.T3", "BRUH Review Guide", accent, label="PAGE 2")
    add_review_grid(doc, problems[6:], accent)
    add_challenge(doc, "Write one mixed decimal application that requires two operations, solve it, and create a plausible first-error example for another student to repair.")
    student = save(doc, "M1.T3-BRUH-Review-Guide-Student.docx")

    key = configure_document(accent, teacher=True)
    add_title(key, "M1.T3", "BRUH Review Guide", accent, label="TEACHER KEY")
    items = [
        KeyItem("1", "0.407, 0.47, 0.704, 0.74", "Compare aligned thousandths."),
        KeyItem("2", "0.63", "The total span is 0.10, so each interval is 0.01."),
        KeyItem("3", "7.078", "Estimate about 7.1; align decimal points."),
        KeyItem("4", "9.33", "Estimate about 9; 12.80 - 3.47 = 9.33."),
        KeyItem("5", "0.72", "The product is less than 1.8 because 0.4 is less than 1."),
        KeyItem("6", "3.12", "2.4(1 + 0.3) = 2.4 + 0.72."),
        KeyItem("7", "609", "Zero holds the tens place; 609 x 4 = 2,436."),
        KeyItem("8", "24", "7.2 ÷ 0.3 = 72 ÷ 3 = 24."),
        KeyItem("9", "5.086", "Align decimal points: 4.700 + 0.386."),
        KeyItem("10", "21 m²", "8.4 x 2.5 = 21; area uses square meters."),
        KeyItem("11", "9 buses", "193 ÷ 24 = 8 R1, so one more bus is required."),
        KeyItem("12", "8", "Both numbers are multiplied by 10, preserving the quotient."),
    ]
    add_answer_table(key, items, accent)
    paragraph(key, "Optional challenge answer", size=10, bold=True, color="7A5600", before=6, after=2)
    paragraph(key, "Answers vary. Require a coherent context, two correct decimal operations, a labeled answer, and a first-error example whose earliest incorrect decision can be identified and repaired. Record challenge points separately.", size=9.5)
    teacher = save(key, "M1.T3-BRUH-Review-Guide-Teacher-Key.docx")
    return student, teacher


def add_assessment_page_one(doc: Document, form: str, accent: str, items: Sequence[Problem]):
    add_title(doc, "M1.T3", f"Topic Assessment - Form {form}", accent, label="20 CORE POINTS")
    add_directions_band(doc, "Section 1 is closed-reference. Complete Items 1-8 without notes. Show work in each box.", accent, title="SECTION 1")
    table = doc.add_table(rows=4, cols=2)
    set_table_geometry(table, (3.75, 3.75))
    for cell, item in zip((c for row in table.rows for c in row.cells), items):
        set_cell_border(cell, LINE, 8)
        set_cell_margins(cell, 95, 115, 95, 115)
        clear_paragraph(cell.paragraphs[0])
        set_run_font(cell.paragraphs[0].add_run(f"{item.number}. {item.prompt}"), 9.2, True, INK)
        for _ in range(3):
            p = cell.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            set_run_font(p.add_run("____________________________________"), 8, False, LINE)
        set_row_min_height(cell._tc.getparent().getparent() if False else table.rows[0], 0.1)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    for row in table.rows:
        prevent_row_split(row)
        set_row_min_height(row, 1.62)


def add_assessment_page_two(doc: Document, form: str, accent: str, items: Sequence[Problem], challenge: str):
    doc.add_page_break()
    add_title(doc, "M1.T3", f"Topic Assessment - Form {form}", accent, label="SECTION 2")
    add_directions_band(doc, "You may use the common one-page reference. Each item is worth 3 points: setup, accurate work, and explanation in context.", accent, title="SECTION 2")
    table = doc.add_table(rows=2, cols=2)
    set_table_geometry(table, (3.75, 3.75))
    for cell, item in zip((c for row in table.rows for c in row.cells), items):
        set_cell_border(cell, LINE, 8)
        set_cell_margins(cell, 95, 115, 85, 115)
        set_cell_shading(cell, WHITE)
        clear_paragraph(cell.paragraphs[0])
        set_run_font(cell.paragraphs[0].add_run(f"{item.number}. {item.prompt}"), 9.1, True, INK)
        if item.labels:
            labels = cell.add_paragraph()
            labels.paragraph_format.space_before = Pt(3)
            labels.paragraph_format.space_after = Pt(2)
            set_run_font(labels.add_run("  |  ".join(item.labels)), 7.8, True, accent)
        for _ in range(5):
            line = cell.add_paragraph()
            line.paragraph_format.space_after = Pt(1)
            set_run_font(line.add_run("____________________________________"), 8, False, LINE)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    for row in table.rows:
        prevent_row_split(row)
        set_row_min_height(row, 2.95)
    paragraph(doc, "", size=1, after=2)
    add_challenge(doc, challenge, lines=1)


def build_assessment(form: str, items_one: Sequence[Problem], items_two: Sequence[Problem], challenge: str,
                     answers: Sequence[KeyItem], challenge_answer: str):
    accent = "7A3341"
    doc = configure_document(accent)
    add_assessment_page_one(doc, form, accent, items_one)
    add_assessment_page_two(doc, form, accent, items_two, challenge)
    student = save(doc, f"M1.T3-Topic-Assessment-Form-{form}-Student.docx")

    key = configure_document(accent, teacher=True)
    add_title(key, "M1.T3", f"Topic Assessment - Form {form}", accent, label="TEACHER KEY AND BLUEPRINT")
    add_directions_band(key, "Core score: 20 points. Items 1-8 are 1 point each. Items 9-12 are 3 points each. Record the optional challenge separately.", accent, title="SCORING")
    add_answer_table(key, answers[:8], accent)
    key.add_page_break()
    add_title(key, "M1.T3", f"Topic Assessment - Form {form}", accent, label="APPLICATION SCORING")
    add_answer_table(key, answers[8:], accent)
    paragraph(key, "Items 9-12 rubric", size=10.5, bold=True, color=accent, before=6, after=2)
    paragraph(key, "3 points: correct setup or representation, accurate computation, and a clear explanation with units or contextual interpretation. 2 points: one element is incomplete. 1 point: meaningful progress. 0 points: no relevant evidence.", size=9.5)
    paragraph(key, "Optional challenge", size=10, bold=True, color="7A5600", before=6, after=2)
    paragraph(key, challenge_answer, size=9.5)
    teacher = save(key, f"M1.T3-Topic-Assessment-Form-{form}-Teacher-Key.docx")
    return student, teacher


def main():
    files: list[Path] = []

    files.extend(build_division_packet(
        "M1.T3.L4-D1",
        "Partial Quotients Bridge",
        "375F4B",
        "Complete all four required problems. Estimate first, connect partial quotients to quotient digits, and verify every quotient.",
        [
            Problem("1", "Represent 864 ÷ 3 with partial quotients and standard division.", ("Estimate", "Partial quotients", "Standard record", "Multiplication check")),
            Problem("2", "A partial-quotient solution for 1,344 ÷ 6 subtracts 1,200, then 120, then 24. Write the corresponding partial quotients and connect them to the digits in the standard quotient.", ("Estimate", "Expanded partial quotients", "Standard record with arrows", "Check")),
            Problem("3", "Solve 756 ÷ 3 with the standard algorithm and verify.", ("Estimate", "Standard record", "Verification")),
            Problem("4", "In 848 ÷ 4 = 212, what does the first quotient digit represent? Explain and verify.", ("Place-value explanation", "Multiplication check")),
        ],
        "Solve 16,536 ÷ 24 with partial quotients and the standard algorithm. Compare which record is more efficient and which makes place value easier to see.",
        [
            KeyItem("1", "288. One useful decomposition is 200 + 80 + 8. Check: 288 x 3 = 864.", "The partial quotients must connect to the hundreds, tens, and ones digits."),
            KeyItem("2", "224. Partial quotients 200 + 20 + 4. Check: 224 x 6 = 1,344.", "Listen for quotient-digit value, not only digit names."),
            KeyItem("3", "252. Check: 252 x 3 = 756.", "First incorrect line identifies whether the need is facts, place value, or subtraction."),
            KeyItem("4", "The 2 represents 200. Check: 212 x 4 = 848.", "The student should name the hundreds place and connect it to 800 ÷ 4 = 200."),
        ],
        "689. Partial quotients could be 600 + 80 + 9. Check: 689 x 24 = 16,536. Accept different valid partial quotients and a reasoned comparison.",
    ))

    files.extend(build_division_packet(
        "M1.T3.L4-D2",
        "Zero Holds the Place",
        "4F6A51",
        "Complete all four required problems. Label quotient places, write any needed zero, and interpret each remainder in the situation.",
        [
            Problem("1", "Solve 2,436 ÷ 4. Explain why the quotient needs a zero.", ("Estimate", "Labeled standard record", "Zero explanation", "Check")),
            Problem("2", "Solve 1,906 ÷ 24. Write the answer as a whole-number quotient with a remainder, then verify with divisor x quotient + remainder = dividend.", ("Estimate", "Standard record", "Remainder check")),
            Problem("3", "A school needs buses for 158 students. Each bus holds 24 students. How many buses are needed?", ("Task Attack notes", "Division work", "Answer in context")),
            Problem("4", "Repair this reasoning: '4 goes into 4 once. Since 4 does not go into the next 1, skip it. Combine the next digits to make 12 and write 3. Then write 2 because 4 goes into 8 twice. Therefore, 4,128 ÷ 4 = 132.' Identify the first incorrect statement and write the correct solution.", ("First incorrect line", "Corrected record", "Why the zero matters")),
        ],
        "Solve 20,005 ÷ 25. Explain both zero placeholders and verify the quotient and remainder.",
        [
            KeyItem("1", "609. The zero holds the tens place because 3 tens cannot be divided into 4 whole groups. 609 x 4 = 2,436.", "A zero at an intermediate place is different from stopping the algorithm."),
            KeyItem("2", "79 R10. Check: 24 x 79 + 10 = 1,906.", "Remainder 10 is valid because it is less than 24."),
            KeyItem("3", "7 buses. 158 ÷ 24 = 6 R14; all students require one additional bus.", "The context changes the reported answer from 6 R14 to 7 buses."),
            KeyItem("4", "1,032. The original work omitted the zero in the hundreds place. 1,032 x 4 = 4,128.", "Have students label quotient places before repairing arithmetic."),
        ],
        "800 R5. The tens and ones quotient places both require zero. Check: 25 x 800 + 5 = 20,005.",
    ))

    files.extend(build_division_packet(
        "M1.T3.L4-D3",
        "Equivalent Decimal Quotients",
        "287A7D",
        "Complete all four required problems. Estimate, write an equivalent division expression, divide, and verify with the original divisor.",
        [
            Problem("1", "Solve 7.35 ÷ 3 and verify.", ("Estimate", "Division record", "Original-expression check")),
            Problem("2", "Solve 7.2 ÷ 0.3 by writing an equivalent whole-divisor expression.", ("Estimate", "Equivalent expression", "Division work", "Check")),
            Problem("3", "Solve 58.8 ÷ 2.4 and check with multiplication.", ("Estimate", "Equivalent expression", "Division work", "Check")),
            Problem("4", "Repair a solution that changes 4.8 ÷ 0.6 into 4.8 ÷ 6. Explain the error.", ("First incorrect line", "Correct equivalent expression", "Correct quotient and check")),
        ],
        "Create a division expression with a divisor less than 1 written to exactly two decimal places and quotient 24. Write two equivalent forms and verify all three.",
        [
            KeyItem("1", "2.45. Check: 2.45 x 3 = 7.35.", "Decimal placement should agree with an estimate near 2.5."),
            KeyItem("2", "7.2 ÷ 0.3 = 72 ÷ 3 = 24. Check: 24 x 0.3 = 7.2.", "Both numbers must be multiplied by 10."),
            KeyItem("3", "58.8 ÷ 2.4 = 588 ÷ 24 = 24.5. Check: 24.5 x 2.4 = 58.8.", "The quotient should be near 60 ÷ 2.5, or about 24."),
            KeyItem("4", "Correct form: 48 ÷ 6 = 8. The original work scaled only the divisor. Check: 8 x 0.6 = 4.8.", "Require parallel scaling arrows on dividend and divisor."),
        ],
        "Answers vary. Example: 6 ÷ 0.25 = 60 ÷ 2.5 = 600 ÷ 25 = 24. Each form must scale both numbers by the same factor.",
    ))

    files.extend(build_division_packet(
        "M1.T3.L4-D4",
        "Application and Error Repair",
        "495A84",
        "Complete all four required problems with Task Attack. Show the operation, interpret the answer, and decide whether it is reasonable.",
        [
            Problem("1", "A display team has 1,906 tiles. Each crate holds 24 tiles. The tiles are blue. Find the number of full crates and tiles remaining, then decide how many crates are needed if every tile must ship.", ("Task Attack notes", "Division work", "Full crates, remainder, and crates needed", "Verification")),
            Problem("2", "A rectangular display has an area of 72.45 square feet and a height of 3.5 feet. Its border is gold. Find its width.", ("Task Attack notes", "Estimate and equation with units", "Equivalent division expression", "Solution and check")),
            Problem("3", "A student begins Problem 2 with 72.45 ÷ 3.5 = 724.5 ÷ 3.5 = 207. Identify the first incorrect equality, then repair the equivalent expression and quotient.", ("First incorrect line", "Correct equivalent expression", "Correct quotient and verification")),
            Problem("4", "A club packs 1,026 cans into boxes that hold 24 cans. Maya calculates 42 R18 and reports that 42 boxes are needed. Decide whether her calculation and conclusion are correct, then repair anything that is incorrect.", ("Check the calculation", "Check the conclusion", "Revised answer in context")),
        ],
        "Create two contexts for 158 ÷ 24: one in which the useful answer is six full groups with 14 remaining, and one in which seven groups or containers are required. Explain why.",
        [
            KeyItem("1", "1,906 ÷ 24 = 79 R10. That means 79 full crates, 10 tiles remaining, and 80 crates needed to ship every tile. Check: 24 x 79 + 10 = 1,906.", "The color is irrelevant. Students must distinguish full crates, remainder, and total containers needed."),
            KeyItem("2", "72.45 ÷ 3.5 = 724.5 ÷ 35 = 20.7 ft. Check: 20.7 x 3.5 = 72.45. A reasonable estimate is about 20 or 21 ft.", "The border color is irrelevant. Preserve square-foot and foot units while connecting area to division."),
            KeyItem("3", "The first equality is incorrect because only the dividend was multiplied by 10. Correct repair: 72.45 ÷ 3.5 = 724.5 ÷ 35 = 20.7.", "Equivalent quotients require the same scale factor for both numbers."),
            KeyItem("4", "The calculation is correct: 1,026 ÷ 24 = 42 R18. The conclusion is incorrect; the remaining 18 cans require one more box, so 43 boxes are needed.", "Separate calculation accuracy from the contextual conclusion."),
        ],
        "Answers vary. Example 1: complete 24-cookie trays gives 6 trays with 14 cookies remaining. Example 2: 24-seat vans for 158 students requires 7 vans. The calculation is the same; the question determines how the remainder is reported.",
    ))

    files.append(build_reference())
    files.extend(build_review())

    form_a_one = [
        Problem("1", "Order 0.075, 0.507, 0.57, and 0.705 from least to greatest."),
        Problem("2", "A line from 1.20 to 1.30 is divided into 5 equal intervals. What value is three intervals to the right of 1.20?"),
        Problem("3", "Find 4.70 + 0.386."),
        Problem("4", "Find 7.40 - 2.685."),
        Problem("5", "Find 2.4 x 1.3."),
        Problem("6", "Find 0.4 x 0.7."),
        Problem("7", "Find 4,128 ÷ 4."),
        Problem("8", "Find 58.8 ÷ 2.4."),
    ]
    form_a_two = [
        Problem("9", "A student writes 4.7 + 0.386 = 0.856 after aligning the final digits. Explain the error and find the correct sum.", ("Explanation", "Correct aligned work", "Correct sum")),
        Problem("10", "A rectangular garden measures 7.5 m by 2.4 m. Estimate and find its exact area. Label the unit.", ("Estimate", "Computation or model", "Answer with unit")),
        Problem("11", "A school is transporting 158 students in buses that hold 24 students. Find the number of buses needed and explain the remainder.", ("Division work", "Remainder interpretation", "Answer in context")),
        Problem("12", "Are 15.36 ÷ 0.48 and 153.6 ÷ 4.8 equivalent? Find the quotient and justify your answer.", ("Scaling relationship", "Quotient", "Justification")),
    ]
    form_a_answers = [
        KeyItem("1", "0.075, 0.507, 0.57, 0.705", "1 point. Decimal comparison and ordering; 6.NS.7 / prerequisite 5.NBT.3b."),
        KeyItem("2", "1.26", "1 point. Scale is 0.02 per interval; number-line reasoning."),
        KeyItem("3", "5.086", "1 point. Add decimals by place value; 6.NS.3."),
        KeyItem("4", "4.715", "1 point. Subtract decimals by place value; 6.NS.3."),
        KeyItem("5", "3.12", "1 point. Decimal multiplication; 6.NS.3."),
        KeyItem("6", "0.28", "1 point. Magnitude and decimal multiplication; 6.NS.3."),
        KeyItem("7", "1,032", "1 point. Multi-digit whole-number division with zero placeholder; 6.NS.2."),
        KeyItem("8", "24.5", "1 point. Decimal division using an equivalent whole divisor; 6.NS.3."),
        KeyItem("9", "Align decimal points: 4.700 + 0.386 = 5.086.", "3 points. Error analysis, place value, and correction; 6.NS.3."),
        KeyItem("10", "Estimate about 18 m²; exact area 18 m².", "3 points. Decimal multiplication in an area context; 6.NS.3 and 6.G.1."),
        KeyItem("11", "7 buses. 158 ÷ 24 = 6 R14, and one more bus is required.", "3 points. Whole-number division and remainder interpretation; 6.NS.2."),
        KeyItem("12", "Yes. Both numbers are multiplied by 10; both quotients are 32.", "3 points. Equivalent quotients and decimal division justification; 6.NS.3."),
    ]
    files.extend(build_assessment(
        "A", form_a_one, form_a_two,
        "After all twelve required items are complete, create and solve a two-operation decimal application, then verify it with a second method.",
        form_a_answers,
        "Answers vary. Require a coherent context, two correct decimal operations, a labeled answer, and a valid second-method verification. Record challenge points separately from the 20-point core score.",
    ))

    form_b_one = [
        Problem("1", "Order 0.086, 0.608, 0.68, and 0.806 from least to greatest."),
        Problem("2", "A line from 2.40 to 2.50 is divided into 5 equal intervals. What value is three intervals to the right of 2.40?"),
        Problem("3", "Find 3.80 + 0.467."),
        Problem("4", "Find 9.20 - 4.586."),
        Problem("5", "Find 1.6 x 2.5."),
        Problem("6", "Find 0.3 x 0.8."),
        Problem("7", "Find 6,024 ÷ 6."),
        Problem("8", "Find 75.6 ÷ 1.8."),
    ]
    form_b_two = [
        Problem("9", "A student writes 3.8 + 0.467 = 0.847 after aligning the final digits. Explain the error and find the correct sum.", ("Explanation", "Correct aligned work", "Correct sum")),
        Problem("10", "A rectangular banner measures 6.4 m by 2.5 m. Estimate and find its exact area. Label the unit.", ("Estimate", "Computation or model", "Answer with unit")),
        Problem("11", "A school is transporting 173 students in buses that hold 28 students. Find the number of buses needed and explain the remainder.", ("Division work", "Remainder interpretation", "Answer in context")),
        Problem("12", "Are 12.6 ÷ 0.6 and 126 ÷ 6 equivalent? Find the quotient and justify your answer.", ("Scaling relationship", "Quotient", "Justification")),
    ]
    form_b_answers = [
        KeyItem("1", "0.086, 0.608, 0.68, 0.806", "1 point. Decimal comparison and ordering; 6.NS.7 / prerequisite 5.NBT.3b."),
        KeyItem("2", "2.46", "1 point. Scale is 0.02 per interval; number-line reasoning."),
        KeyItem("3", "4.267", "1 point. Add decimals by place value; 6.NS.3."),
        KeyItem("4", "4.614", "1 point. Subtract decimals by place value; 6.NS.3."),
        KeyItem("5", "4", "1 point. Decimal multiplication; 6.NS.3."),
        KeyItem("6", "0.24", "1 point. Magnitude and decimal multiplication; 6.NS.3."),
        KeyItem("7", "1,004", "1 point. Multi-digit whole-number division with zero placeholder; 6.NS.2."),
        KeyItem("8", "42", "1 point. Decimal division using an equivalent whole divisor; 6.NS.3."),
        KeyItem("9", "Align decimal points: 3.800 + 0.467 = 4.267.", "3 points. Error analysis, place value, and correction; 6.NS.3."),
        KeyItem("10", "Estimate about 16 m²; exact area 16 m².", "3 points. Decimal multiplication in an area context; 6.NS.3 and 6.G.1."),
        KeyItem("11", "7 buses. 173 ÷ 28 = 6 R5, and one more bus is required.", "3 points. Whole-number division and remainder interpretation; 6.NS.2."),
        KeyItem("12", "Yes. Both numbers are multiplied by 10; both quotients are 21.", "3 points. Equivalent quotients and decimal division justification; 6.NS.3."),
    ]
    files.extend(build_assessment(
        "B", form_b_one, form_b_two,
        "After all twelve required items are complete, create and solve a two-operation decimal application, then verify it with a second method.",
        form_b_answers,
        "Answers vary. Require a coherent context, two correct decimal operations, a labeled answer, and a valid second-method verification. Record challenge points separately from the 20-point core score.",
    ))

    for path in files:
        print(path)


if __name__ == "__main__":
    main()
